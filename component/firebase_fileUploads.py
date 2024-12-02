import firebase_admin
from firebase_admin import credentials, storage, firestore
from google.cloud.firestore_v1.base_query import FieldFilter, Or
import json
import os
import io
from PyPDF2 import PdfReader

from langchain.docstore.document import Document
from docx import Document as DocxDocument



try:
    if not os.path.exists('cred_files'):
        os.makedirs('cred_files')

    cred_file_path=r'./cred_files/commission-83ab7-firebase-adminsdk-4dzjr-fc87f36497.json'
    # cred_file_path= r'./uploads/smartcard-85e1b-firebase-adminsdk-demdi-51036c239b.json'

    # Load the service account key from an environment variable
    with open(cred_file_path, 'r') as file:
        service_account_key_json = json.load(file)

except:
    if not os.path.exists('cred_files'):
        os.makedirs('cred_files')
    cred_file_path=r'../cred_files/commission-83ab7-firebase-adminsdk-4dzjr-fc87f36497.json'
    # cred_file_path= r'./uploads/smartcard-85e1b-firebase-adminsdk-demdi-51036c239b.json'

    # Load the service account key from an environment variable
    with open(cred_file_path, 'r') as file:
        service_account_key_json = json.load(file)

if service_account_key_json is None:
    raise ValueError("FIREBASE_SERVICE_ACCOUNT_KEY environment variable is not set")

# Parse the JSON string into a dictionary
try:
    service_account_info = service_account_key_json
except json.JSONDecodeError:
    raise ValueError("FIREBASE_SERVICE_ACCOUNT_KEY is not a valid JSON string")

# Initialize Firebase app
cred = credentials.Certificate(service_account_info)
app = firebase_admin.initialize_app(cred, {
    'storageBucket': os.getenv('STORAGE_BUCKET')
})


bucket= storage.bucket(app= app)





def create_folder_upload_files(folder_path, file):
    try: 
        blob= bucket.blob(folder_path)
        blob.upload_from_file(file.file, content_type=file.content_type)
        file.file.seek(0)
        return 'file uploaded'
    except Exception as e:
        raise e


def retrieve_collection_name_from_firebase(path):
    try:
        path= f'users/{path}'
        blobs= bucket.list_blobs(prefix= path)
        
        result=set()
        for blob in blobs:
            # arr= np.frombuffer(blob.download_as_string(), np.uint8)

            result.add(blob.name.split('/')[-2])
            
            # print(blob.name.split('/')[:-1])
            # print('hello')

        return result
            # break
        
            # print(arr)
    except Exception as e:
        print('why')
        raise e
   
def retrieve_collection_from_firebase(folder_path):
    try:
        parts = folder_path.split('/')
        #test_path=f'users/folders/{folder_path}'
        # users/{id}/folders/{folder_names}/{files}
        folder_path= f"users/{parts[0]}/folders/{parts[1]}/"

        print

        blobs= bucket.list_blobs(prefix= folder_path)
        
        result=[]
        for blob in blobs:
            # print(f'{folder_path}_.pdf')
            # arr= np.frombuffer(blob.download_as_string(), np.uint8)
            if blob.name==f'{folder_path}_.pdf':
                continue
            content= blob.download_as_bytes()
            file_type= blob.content_type
            
            if blob.content_type=='application/pdf':
                
                pdf_file= io.BytesIO(content)
                reader= PdfReader(pdf_file) 
                #pdf_text=''
                for pnum, page in enumerate(reader.pages, start=1):
                    pdf_text = page.extract_text()

                    doc = Document(page_content=pdf_text, metadata={"filename": blob.name, 'page number':pnum})
                    result.append(doc)
                

            elif blob.content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                # Process DOCX file
                docx_file = io.BytesIO(content)
                docx_document = DocxDocument(docx_file)
                docx_text = ''
                for para in docx_document.paragraphs:
                    if para.text:
                        docx_text += para.text + '\n'

                
                doc = Document(page_content=docx_text, metadata={"filename": blob.name})
                result.append(doc)
                

            elif blob.content_type == 'text/plain':
                # Process TXT file
                txt_content = content.decode('utf-8')
                
                doc = Document(page_content=txt_content, metadata={"filename": blob.name})
                result.append(doc)


            elif blob.content_type=='application/json':
                reader= content.decode('utf-8')
                json_file= json.loads(reader)
                path= blob.name
                parts = path.split('/')
                data= {parts[2]:json_file}
                result.append(data)

            elif blob.content_type == 'text/csv':
                #content = blob.download_as_bytes()
                csv_content = content.decode('utf-8')
                result.append(csv_content)

            elif blob.content_type in  ['application/vnd.ms-excel', 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet']: # for xls, xlsx
                #content = blob.download_as_bytes()
                print("type correct")
                result.append(content)

            elif blob.content_type == 'text/xml':
                # Process TXT file
                xml_content = content.decode('utf-8')
                
                # doc = Document(page_content=txt_content, metadata={"filename": blob.name})
                result.append(xml_content)

            elif blob.content_type in ['application/octet-stream', 'application/x-subrip', 'text/srt']:
                # Process TXT file
                srt_content = content.decode('utf-8')
                
                # doc = Document(page_content=txt_content, metadata={"filename": blob.name})
                result.append(srt_content)
            

        return result, file_type
            
    except Exception as e:
        raise e


def retrieve_file_from_firebase(folder_path):
    try:
        parts = folder_path.split('/')
        
        file_path= f"users/{parts[0]}/folders/{parts[1]}/"

        file_name= parts[2]

        blobs= bucket.list_blobs(prefix= file_path)
        

        
        result=[]
        for blob in blobs:
            
            if blob.name==f'{folder_path}_.pdf':
                continue
            content= blob.download_as_bytes()
            file_type= blob.content_type
            
            if blob.content_type=='application/pdf':
                
                if blob.name== f"{file_path}{file_name}":
                    return blob.name, file_type
                

            elif blob.content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                # Process DOCX file
                if blob.name== f"{file_path}{file_name}":
                    return blob.name, file_type
                

            elif blob.content_type == 'text/plain':
                # Process TXT file
                if blob.name== f"{file_path}{file_name}":
                    return blob.name, file_type

            elif blob.content_type=='application/json':
                if blob.name== f"{file_path}{file_name}":
                    reader= content.decode('utf-8')
                    json_file= json.loads(reader)
                    path= blob.name
                    parts = path.split('/')
                    data= {parts[2]:json_file}
                    result.append(data)

            elif blob.content_type == 'text/csv':
                if blob.name== f"{file_path}{file_name}":
                    csv_content = content.decode('utf-8')
                    result.append(csv_content)

            elif blob.content_type in  ['application/vnd.ms-excel', 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet']: # for xls, xlsx
                if blob.name== f"{file_path}{file_name}":
                    print(f"{file_path}{file_name}")
                    print(f"blob.name: {blob.name}")
                    print("type correct")
                    result.append(content)

            elif blob.content_type == 'text/xml':
                if blob.name== f"{file_path}{file_name}":
                    print(f"{file_path}{file_name}")
                    print(f"blob.name: {blob.name}")
                    xml_content = content.decode('utf-8')
                    
                    # doc = Document(page_content=txt_content, metadata={"filename": blob.name})
                    result.append(xml_content)

            elif blob.content_type in ['application/octet-stream', 'application/x-subrip', 'text/srt']:
                if blob.name== f"{file_path}{file_name}":
                    srt_content = content.decode('utf-8')
                    
                    # doc = Document(page_content=txt_content, metadata={"filename": blob.name})
                    result.append(srt_content)
            

        return result, file_type
            # break
        
            # print(arr)
    except Exception as e:
        raise e


